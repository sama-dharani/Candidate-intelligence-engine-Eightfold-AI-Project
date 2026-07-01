import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Union

try:
    import docx
    _DOCX_BACKEND = "python-docx"
except ImportError:
    _DOCX_BACKEND = "zip"

class DOCXLoader:
    """Extract raw text from DOCX files.
    
    Uses python-docx if available, otherwise falls back to parsing the underlying XML via zipfile.
    """
    
    def load(self, file_path: Union[str, Path]) -> str:
        path = Path(file_path)
        if not path.is_file():
            raise FileNotFoundError(f"File not found: {path}")

        # Try python-docx
        if _DOCX_BACKEND == "python-docx":
            try:
                doc = docx.Document(str(path))
                text = []
                for para in doc.paragraphs:
                    text.append(para.text)
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        text.append("\t".join(row_text))
                return "\n".join(text)
            except Exception:
                pass
                
        # Fallback: parse XML from zip
        try:
            with zipfile.ZipFile(path) as docx_zip:
                xml_content = docx_zip.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            
            # The XML namespaces usually used by docx
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            paragraphs = []
            for p in tree.findall('.//w:p', namespaces):
                texts = [node.text for node in p.findall('.//w:t', namespaces) if node.text]
                if texts:
                    paragraphs.append(''.join(texts))
            return "\n".join(paragraphs)
        except Exception:
            return ""
